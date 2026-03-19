from docx.api import Document
from docx.shared import Cm
from docx.shared import Pt



def writerbold(data, name):
    if name == "":
        name = "demo.docx"
    else:
        name = name+".docx"
    
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Arial Narrow'
    style.font.size = Pt(12)
    i = 4
    counter = 1
    for n in data:
        if i == 4 :
          #  num = document.add_paragraph(str(counter)+"."+"\n")
            counter += 1
        i = i-1
        p = document.add_paragraph("")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.left_indent = Cm(-2)
        p.add_run(n[0]).bold = True
        p.add_run(n[1]).bold = True
        p.add_run(n[2])
        if i == 0:
            i = 4 

    document.add_page_break()
    document.save('demo.docx')


#############################
def reader(path):
    document = Document(path)
    Text=[]
    for p in document.paragraphs:
        p1 = p.text
        if p1 != "":
            if p1[0] == "A" or p1[0] == "B" or p1[0] == "C" or p1[0] == "D":
                if p1[1] == ".":
                    p2 = p1[2:]
                    Text.append([p1[0],p1[1],p2])
    return Text